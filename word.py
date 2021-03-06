import io
import json
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from sphinx.util import requests
from docxtpl import DocxTemplate, RichText


class wordSave:
    def __init__(self):
        self.makeTitul()
        self.document = Document()
        self.helpDict = {
            "paragraph": self.addLine,
            "heading": self.addHeading,
            "list": self.addList,
            "link": self.addLink
        }

    def addPageBreak(self):
        self.document.add_page_break()

    def addCode(self, code):
        par = self.document.add_paragraph("")
        for line in code:
            par.add_run(line).font.size = Pt(10)

    def makeTitul(self):
        tpl = DocxTemplate('patternTitle.docx')
        self.js = json.load(open('settings.json'))
        context = {
            'cathedra': RichText(self.js['cathedra'].encode('cp1251', 'ignore')),
            'discipline': RichText(self.js['discipline'].encode('cp1251', 'ignore')),
            'theme': RichText(self.js['theme'].encode('cp1251', 'ignore')),
            'group': RichText(self.js['group'].encode('cp1251', 'ignore')),
            'name': RichText(self.js['name'].encode('cp1251', 'ignore')),
            'teacher': RichText(self.js['teacher'].encode('cp1251', 'ignore')),
            'initialData': RichText(self.js['initialData'].encode('cp1251', 'ignore')),
            'content': RichText(self.js['content'].encode('cp1251', 'ignore')),
            'numberPages ': RichText(self.js['numberPages'].encode('cp1251', 'ignore')),
            'surrender': RichText(self.js['surrender'].encode('cp1251', 'ignore')),
            'numberPages': RichText(self.js['numberPages'].encode('cp1251', 'ignore')),
            'extradition': RichText(self.js['extradition'].encode('cp1251', 'ignore')),
            'protection': RichText(self.js['protection'].encode('cp1251', 'ignore')),
            'annotation': RichText(self.js['annotation'].encode('cp1251', 'ignore')),
            'summary': RichText(self.js['summary'].encode('cp1251', 'ignore')),
            'introduction': RichText(self.js['introduction'].encode('cp1251', 'ignore')),
        }
        tpl.render(context)
        tpl.save('compileTitul.docx')


    def addLine(self, line):
        self.document.add_paragraph(line)

    def addHeadCode(self, line):
        self.document.add_heading().add_run(line).font.size = Pt(20)

    # def addFromDict(self,dict):

    def reform(self, mass, line, tag):
        # line.style=self.document.styles['Normal']
        for element in mass:
            x = line.add_run(element["line"])
            if tag in self.js['formatter']:
                x.font.size = Pt(self.js['formatter'][tag]['size'])
            x.font.bold = element["bold"]
            x.font.italic = element["italic"]
            x.font.underline = element["underline"]

    def addHeading(self, mass, lvl=1):
        line = self.document.add_heading('', level=lvl)
        self.reform(mass, line, "h" + str(lvl))

    # style='IntenseQuote'-цитата
    # style='ListBullet'-точки
    # style='ListNumber'-цифры

    def addList(self, mass):
        line = self.document.add_paragraph('', style='ListBullet')
        self.reform(mass, line, "list")

    def checkInnerInfo(self, mass, line):
        for elem in mass:
            if elem['type'] == "img":
                self.addPicture([elem])
            if elem['type'] == "link":
                if elem['line'] == "":
                    self.add_hyperlink(line, elem['src'], " [url]")
                else:
                    self.add_hyperlink(line, elem['src'], elem['line'])

    def addLink(self, mass):
        line = self.document.add_paragraph('')
        self.checkInnerInfo(mass, line)

    def add_hyperlink(self, paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def addParagraph(self, mass):
        line = self.document.add_paragraph('')
        self.reform(mass, line, "p")
        self.checkInnerInfo(mass, line)

    def getImgByUrl(self, url):
        response = requests.get(url, stream=True)
        return io.BytesIO(response.content)

    def addPicture(self, mass):
        for elem in mass:
            if elem['type'] == 'img':
                url = elem["src"]
                try:
                    self.document.add_picture(self.getImgByUrl(url))
                except Exception:
                    print("bad image url")
            else:
                if elem['type'] == 'link':
                    self.addLink([elem])
                else:
                    self.addParagraph([elem])

    def saveFile(self):
        self.document.save('ReadyProject.docx')
